"""
Tinku Phase 6 — Export Router
Report PDF: 9.5/10 | Report DOCX: 8.8/10 | WPS Android safe
"""
from fastapi import APIRouter
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
import io
import re

router = APIRouter()

class ExportRequest(BaseModel):
    content: str
    format: str   # "pdf" or "docx"
    title: str = "Tinku Report"

def extract_title(content: str) -> str:
    for line in content.strip().split('\n')[:5]:
        clean = re.sub(r'[#*`]', '', line).strip()
        if 5 < len(clean) < 100:
            return clean
    return "Tinku Report"

def strip_emoji(text: str) -> str:
    return re.sub(r'[\U0001F000-\U0001FFFF]|[\u2600-\u27BF]|\u200d', '', text).strip()

def clean_inline(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*',     r'<i>\1</i>', text)
    text = re.sub(r'`(.+?)`',
        r'<font name="Courier" size="9" color="#4338CA">\1</font>', text)
    return strip_emoji(text)

@router.post("/export")
async def export_content(data: ExportRequest):
    content = data.content
    title   = extract_title(content) if data.title == "Tinku Report" else data.title
    if data.format == "pdf":
        return await generate_pdf(content, title)
    elif data.format == "docx":
        return await generate_docx(content, title)
    return JSONResponse({"error": "Invalid format"}, status_code=400)


# ═══════════════════════════════════════════════════════════
# PDF REPORT GENERATOR — v6 (9.5/10)
# ═══════════════════════════════════════════════════════════
async def generate_pdf(content: str, title: str) -> StreamingResponse:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                         HRFlowable, KeepTogether)
        from reportlab.lib.enums import TA_JUSTIFY
        from reportlab.pdfgen import canvas as pdfcanvas
        from reportlab.platypus.flowables import Flowable

        A4_W, A4_H = A4
        INDENT = 1.8 * cm

        ACCENT       = colors.HexColor('#4338CA')
        ACCENT2      = colors.HexColor('#6366f1')
        ACCENT_LIGHT = colors.HexColor('#EEF2FF')
        ACCENT_MID   = colors.HexColor('#C7D2FE')
        ACCENT_DARK  = colors.HexColor('#3730a3')
        DARK         = colors.HexColor('#0f172a')
        GRAY         = colors.HexColor('#64748b')
        LGRAY        = colors.HexColor('#e2e8f0')
        ROW_BG       = colors.HexColor('#EEF2FF')
        LGRAY2       = colors.HexColor('#f8fafc')
        WHITE        = colors.HexColor('#ffffff')
        BODY_C       = colors.HexColor('#1e293b')
        DOC_W        = A4_W - 2 * INDENT

        class ShadedH1(Flowable):
            def __init__(self, text, doc_width):
                super().__init__()
                self.text = text; self.width = doc_width; self.height = 36
            def wrap(self, aw, ah):
                self.width = aw; return (aw, self.height)
            def draw(self):
                c = self.canv
                c.setFillColor(ACCENT_LIGHT)
                c.roundRect(0, 0, self.width, self.height, 6, fill=1, stroke=0)
                c.setFillColor(ACCENT)
                c.rect(0, 0, 7, self.height, fill=1, stroke=0)
                c.setFillColor(ACCENT_MID)
                c.rect(self.width - 5, 0, 5, self.height, fill=1, stroke=0)
                c.setFillColor(ACCENT)
                c.setFont('Helvetica-Bold', 13)
                c.drawString(18, 12, self.text[:82])

        class NumberedItem(Flowable):
            def __init__(self, num, label, rest, doc_width):
                super().__init__()
                self.num = str(num); self.label = label
                self.rest = rest; self.width = doc_width; self.height = 36
            def _compute_lines(self, c):
                badge_end = 42
                label_w = c.stringWidth(self.label + ':  ', 'Helvetica-Bold', 10.5) if self.label else 0
                avail_w = self.width - badge_end - label_w - 10
                words = self.rest.split()
                line, lines = [], []
                for w in words:
                    test = ' '.join(line + [w])
                    if c.stringWidth(test, 'Helvetica', 10.5) <= avail_w:
                        line.append(w)
                    else:
                        if line: lines.append(' '.join(line))
                        line = [w]
                if line: lines.append(' '.join(line))
                return lines or ['']
            def wrap(self, aw, ah):
                self.width = aw
                n = max(1, len(self.rest) // 60 + 1)
                self.height = max(36, 16 * n + 22)
                return (aw, self.height)
            def draw(self):
                c = self.canv
                lines = self._compute_lines(c)
                self.height = max(36, 16 * len(lines) + 22)
                mid = self.height / 2
                c.setFillColor(ROW_BG)
                c.roundRect(0, 2, self.width, self.height - 4, 6, fill=1, stroke=0)
                c.setFillColor(ACCENT)
                c.roundRect(0, 2, 5, self.height - 4, 3, fill=1, stroke=0)
                badge_w = 24; bx = 10
                c.setFillColor(ACCENT)
                c.roundRect(bx, mid - 11, badge_w, 22, 4, fill=1, stroke=0)
                c.setFillColor(WHITE)
                c.setFont('Helvetica-Bold', 10)
                c.drawCentredString(bx + badge_w / 2, mid - 4, self.num)
                x = 42; top_y = mid + (len(lines) - 1) * 8
                if self.label:
                    c.setFillColor(ACCENT)
                    c.setFont('Helvetica-Bold', 10.5)
                    lw = c.stringWidth(self.label + ':  ', 'Helvetica-Bold', 10.5)
                    c.drawString(x, top_y - 4, self.label + ':')
                    x += lw
                c.setFillColor(BODY_C)
                c.setFont('Helvetica', 10.5)
                for i, ln in enumerate(lines):
                    c.drawString(x if i == 0 else 42, top_y - 4 - i * 16, ln)

        class ReportCanvas(pdfcanvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []
            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()
            def save(self):
                total = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self._draw_chrome(total)
                    super().showPage()
                super().save()
            def _draw_chrome(self, total):
                self.saveState()
                self.setFillColor(ACCENT)
                self.rect(0, A4_H - 1.75*cm, A4_W, 1.75*cm, fill=1, stroke=0)
                self.setFillColor(ACCENT2)
                self.rect(0, A4_H - 1.77*cm, A4_W, 0.08*cm, fill=1, stroke=0)
                self.setFillColor(ACCENT_DARK)
                self.circle(INDENT - 0.5*cm, A4_H - 0.88*cm, 0.22*cm, fill=1, stroke=0)
                self.setFillColor(colors.HexColor('#818cf8'))
                self.circle(INDENT - 0.12*cm, A4_H - 0.88*cm, 0.11*cm, fill=1, stroke=0)
                self.setFillColor(WHITE)
                self.setFont('Helvetica-Bold', 11)
                self.drawString(INDENT + 0.15*cm, A4_H - 1.08*cm, strip_emoji(title))
                pill_w, pill_h = 1.9*cm, 0.58*cm
                px = A4_W - INDENT - pill_w; py = A4_H - 1.32*cm
                self.setFillColor(ACCENT_DARK)
                self.roundRect(px, py, pill_w, pill_h, 5, fill=1, stroke=0)
                self.setFillColor(WHITE)
                self.setFont('Helvetica-Bold', 8)
                self.drawCentredString(px + pill_w / 2, py + 0.17*cm, 'Tinku AI')
                self.setFillColor(LGRAY2)
                self.rect(0, 0, A4_W, 1.05*cm, fill=1, stroke=0)
                self.setFillColor(ACCENT_MID)
                self.rect(0, 1.03*cm, A4_W, 0.06*cm, fill=1, stroke=0)
                self.setFillColor(ACCENT)
                self.circle(INDENT - 0.3*cm, 0.42*cm, 0.07*cm, fill=1, stroke=0)
                self.setFillColor(GRAY)
                self.setFont('Helvetica-Oblique', 7.5)
                self.drawString(INDENT, 0.35*cm, 'Generated by Tinku AI')
                self.setFont('Helvetica', 7.5)
                self.drawCentredString(A4_W / 2, 0.35*cm, '—')
                self.drawRightString(A4_W - INDENT, 0.35*cm,
                                     f'Page {self._pageNumber} of {total}')
                self.restoreState()

        title_s  = ParagraphStyle('T',  fontSize=26, fontName='Helvetica-Bold',
            textColor=DARK, spaceAfter=4, spaceBefore=4, leading=32)
        meta_s   = ParagraphStyle('M',  fontSize=9,  fontName='Helvetica',
            textColor=GRAY, spaceAfter=18)
        h2_s     = ParagraphStyle('H2', fontSize=11, fontName='Helvetica-Bold',
            textColor=DARK, spaceBefore=12, spaceAfter=5, leading=16)
        body_s   = ParagraphStyle('B',  fontSize=10.5, fontName='Helvetica',
            textColor=BODY_C, spaceAfter=9, leading=18,
            alignment=TA_JUSTIFY, wordWrap='LTR',
            allowWidows=0, allowOrphans=0)
        bullet_s = ParagraphStyle('BL', fontSize=10.5, fontName='Helvetica',
            textColor=BODY_C, spaceAfter=6, leading=17, leftIndent=20)

        story = []
        story.append(Spacer(1, 6))
        story.append(Paragraph(strip_emoji(title), title_s))
        story.append(Paragraph('Generated by Tinku AI', meta_s))
        story.append(HRFlowable(width='100%', thickness=3, color=ACCENT,
                                 lineCap='round', spaceAfter=22))

        for line in content.strip().split('\n'):
            s = line.strip()
            if not s:
                story.append(Spacer(1, 4)); continue
            if re.match(r'^#{1,2}\s', s):
                text = strip_emoji(re.sub(r'^#{1,2}\s+|[*`]', '', s)).strip()
                story.append(Spacer(1, 14))
                story.append(ShadedH1(text, DOC_W))
                story.append(Spacer(1, 12))
            elif re.match(r'^#{3,}\s', s):
                text = strip_emoji(re.sub(r'^#{3,}\s+|[*`]', '', s)).strip()
                story.append(Paragraph(text, h2_s))
            elif m := re.match(r'^(\d+)\.\s+(.+)', s):
                num = m.group(1); raw = m.group(2)
                lm  = re.match(r'\*\*(.+?)\*\*[:\s]*(.*)', raw)
                label = strip_emoji(lm.group(1)) if lm else ''
                rest  = strip_emoji(re.sub(r'[*`]', '',
                        lm.group(2) if lm else re.sub(r'\*\*(.+?)\*\*', r'\1', raw)))
                story.append(KeepTogether([NumberedItem(num, label, rest, DOC_W), Spacer(1, 7)]))
            elif re.match(r'^[-*•]\s', s):
                raw = s[2:].strip()
                lm  = re.match(r'\*\*(.+?)\*\*[:\s]*(.*)', raw)
                if lm:
                    label = strip_emoji(lm.group(1))
                    rest  = strip_emoji(re.sub(r'[*`]', '', lm.group(2)))
                    text  = (f'<font color="#6366f1"><b>▸</b></font>  '
                             f'<b><font color="#4338CA">{label}</font></b>: {rest}')
                else:
                    text = f'<font color="#6366f1"><b>▸</b></font>  {clean_inline(raw)}'
                story.append(Paragraph(text, bullet_s))
            elif s in ('---', '***', '___'):
                story.append(HRFlowable(width='100%', thickness=0.5, color=LGRAY, spaceAfter=6))
            else:
                text = clean_inline(s)
                if text:
                    try:    story.append(Paragraph(text, body_s))
                    except: story.append(Paragraph(strip_emoji(re.sub(r'[*#`]', '', s)), body_s))

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
            leftMargin=INDENT, rightMargin=INDENT,
            topMargin=1.95*cm, bottomMargin=1.4*cm,
            title=title, author='Tinku AI')
        doc.build(story, canvasmaker=ReportCanvas)
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="tinku_report.pdf"'})

    except Exception as e:
        import traceback
        return JSONResponse({"error": f"PDF failed: {str(e)}\n{traceback.format_exc()}"}, status_code=500)


# ═══════════════════════════════════════════════════════════
# DOCX REPORT GENERATOR — v6 (8.8/10, WPS Android safe)
# ═══════════════════════════════════════════════════════════
async def generate_docx(content: str, title: str) -> StreamingResponse:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        ACCENT_HEX  = '4338CA'
        ACCENT2_HEX = '6366f1'
        ACCENTL_HEX = 'EEF2FF'
        ROW_HEX     = 'EEF2FF'
        DARK_RGB    = RGBColor(0x0f, 0x17, 0x2a)
        GRAY_RGB    = RGBColor(0x64, 0x74, 0x8b)
        BODY_RGB    = RGBColor(0x1e, 0x29, 0x3b)
        WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)
        ACCENT_RGB  = RGBColor(0x43, 0x38, 0xCA)
        ACCENT2_RGB = RGBColor(0x63, 0x66, 0xf1)

        def set_cell_shd(cell, fill):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill); tcPr.append(shd)

        def set_para_shd(para, fill):
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill); pPr.append(shd)

        def remove_cell_margins(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            mar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:w'), '0'); el.set(qn('w:type'), 'dxa')
                mar.append(el)
            tcPr.append(mar)

        def bottom_border(para, color=ACCENT2_HEX, sz='6'):
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            el   = OxmlElement('w:bottom')
            el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '4');    el.set(qn('w:color'), color)
            pBdr.append(el); pPr.append(pBdr)

        def field_run(para, instr):
            for tag, val in [('begin', None), ('instr', instr), ('end', None)]:
                r = OxmlElement('w:r')
                rpr = OxmlElement('w:rPr')
                sz  = OxmlElement('w:sz');    sz.set(qn('w:val'), '16')
                col = OxmlElement('w:color'); col.set(qn('w:val'), '64748b')
                ita = OxmlElement('w:i')
                rpr.append(sz); rpr.append(col); rpr.append(ita); r.append(rpr)
                if tag == 'instr':
                    it = OxmlElement('w:instrText')
                    it.set(qn('xml:space'), 'preserve'); it.text = val; r.append(it)
                else:
                    fc = OxmlElement('w:fldChar')
                    fc.set(qn('w:fldCharType'), tag); r.append(fc)
                para._p.append(r)

        def add_footer(doc):
            section = doc.sections[0]
            footer  = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_shd(fp, 'F8FAFC')
            r1 = fp.add_run('Generated by Tinku AI    •    Page ')
            r1.font.size = Pt(8); r1.font.color.rgb = GRAY_RGB; r1.font.italic = True
            field_run(fp, ' PAGE ')
            r2 = fp.add_run(' of ')
            r2.font.size = Pt(8); r2.font.color.rgb = GRAY_RGB; r2.font.italic = True
            field_run(fp, ' NUMPAGES ')

        def make_h1_table(doc, text):
            sb = doc.add_paragraph()
            sb.paragraph_format.space_before = Pt(0); sb.paragraph_format.space_after = Pt(8)
            tbl = doc.add_table(rows=1, cols=2)
            tbl.autofit = False
            tbl.columns[0].width = Cm(0.4); tbl.columns[1].width = Cm(16.6)
            set_cell_shd(tbl.cell(0,0), ACCENT_HEX); remove_cell_margins(tbl.cell(0,0))
            tbl.cell(0,0).paragraphs[0].paragraph_format.space_before = Pt(9)
            tbl.cell(0,0).paragraphs[0].paragraph_format.space_after  = Pt(9)
            set_cell_shd(tbl.cell(0,1), ACCENTL_HEX)
            tp = tbl.cell(0,1).paragraphs[0]
            tp.paragraph_format.space_before = Pt(9); tp.paragraph_format.space_after = Pt(9)
            tp.paragraph_format.left_indent  = Cm(0.3)
            r = tp.add_run(text)
            r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = ACCENT_RGB
            sa = doc.add_paragraph()
            sa.paragraph_format.space_before = Pt(0); sa.paragraph_format.space_after = Pt(6)

        def make_numbered_table(doc, num, label, rest):
            tbl = doc.add_table(rows=1, cols=2)
            tbl.autofit = False
            tbl.columns[0].width = Cm(0.75); tbl.columns[1].width = Cm(16.25)
            bc = tbl.cell(0,0); set_cell_shd(bc, ACCENT_HEX); remove_cell_margins(bc)
            bp = bc.paragraphs[0]; bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.paragraph_format.space_before = Pt(8); bp.paragraph_format.space_after = Pt(8)
            br = bp.add_run(num)
            br.font.bold = True; br.font.size = Pt(11); br.font.color.rgb = WHITE_RGB
            tc2 = tbl.cell(0,1); set_cell_shd(tc2, ROW_HEX)
            tp3 = tc2.paragraphs[0]
            tp3.paragraph_format.space_before = Pt(8); tp3.paragraph_format.space_after = Pt(8)
            tp3.paragraph_format.left_indent  = Cm(0.25)
            if label:
                lr2 = tp3.add_run(label + ':  ')
                lr2.font.bold = True; lr2.font.size = Pt(10.5); lr2.font.color.rgb = ACCENT_RGB
            if rest:
                tr2 = tp3.add_run(rest)
                tr2.font.size = Pt(10.5); tr2.font.color.rgb = BODY_RGB
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(4)

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(1.2); section.bottom_margin = Cm(2.2)
            section.left_margin   = Cm(2.0); section.right_margin  = Cm(2.0)

        # Header
        htbl = doc.add_table(rows=1, cols=2)
        htbl.autofit = False
        htbl.columns[0].width = Cm(12); htbl.columns[1].width = Cm(5)
        for cell in htbl.row_cells(0):
            set_cell_shd(cell, ACCENT_HEX); remove_cell_margins(cell)
        lp = htbl.cell(0,0).paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_before = Pt(10); lp.paragraph_format.space_after = Pt(10)
        lp.paragraph_format.left_indent  = Cm(0.4)
        lr = lp.add_run(strip_emoji(title))
        lr.font.bold = True; lr.font.size = Pt(11); lr.font.color.rgb = WHITE_RGB
        rp = htbl.cell(0,1).paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(12); rp.paragraph_format.space_after = Pt(10)
        rp.paragraph_format.right_indent = Cm(0.4)
        rr = rp.add_run('Tinku AI')
        rr.font.size = Pt(9); rr.font.color.rgb = WHITE_RGB; rr.font.bold = True

        # Accent stripe
        stripe = doc.add_table(rows=1, cols=1)
        stripe.autofit = False; stripe.columns[0].width = Cm(17)
        set_cell_shd(stripe.cell(0,0), ACCENT2_HEX); remove_cell_margins(stripe.cell(0,0))
        stripe.cell(0,0).paragraphs[0].paragraph_format.space_before = Pt(2)
        stripe.cell(0,0).paragraphs[0].paragraph_format.space_after  = Pt(2)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Title block
        tp = doc.add_paragraph()
        tr = tp.add_run(strip_emoji(title))
        tr.font.bold = True; tr.font.size = Pt(22); tr.font.color.rgb = DARK_RGB
        tp.paragraph_format.space_after = Pt(3)
        dp = doc.add_paragraph()
        dr = dp.add_run('Generated by Tinku AI')
        dr.font.size = Pt(9); dr.font.color.rgb = GRAY_RGB; dr.font.italic = True
        dp.paragraph_format.space_after = Pt(10)

        # Divider
        div = doc.add_table(rows=1, cols=1)
        div.autofit = False; div.columns[0].width = Cm(17)
        set_cell_shd(div.cell(0,0), ACCENT_HEX); remove_cell_margins(div.cell(0,0))
        div.cell(0,0).paragraphs[0].paragraph_format.space_before = Pt(3)
        div.cell(0,0).paragraphs[0].paragraph_format.space_after  = Pt(3)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Parse content
        for line in content.strip().split('\n'):
            s = line.strip()
            if not s:
                ep = doc.add_paragraph(); ep.paragraph_format.space_after = Pt(2); continue
            if re.match(r'^#{1,2}\s', s):
                make_h1_table(doc, strip_emoji(re.sub(r'^#{1,2}\s+|[*`]', '', s)).strip())
            elif re.match(r'^#{3,}\s', s):
                text = strip_emoji(re.sub(r'^#{3,}\s+|[*`]', '', s)).strip()
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(10); h.paragraph_format.space_after = Pt(5)
                bottom_border(h, ACCENT2_HEX, '6')
                r = h.add_run(text)
                r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK_RGB
            elif m := re.match(r'^(\d+)\.\s+(.+)', s):
                num = m.group(1); raw = m.group(2)
                lm  = re.match(r'\*\*(.+?)\*\*[:\s]*(.*)', raw)
                label = strip_emoji(lm.group(1)) if lm else ''
                rest  = strip_emoji(re.sub(r'[*`]', '',
                        lm.group(2) if lm else re.sub(r'\*\*(.+?)\*\*', r'\1', raw)))
                make_numbered_table(doc, num, label, rest)
            elif re.match(r'^[-*•]\s', s):
                raw = s[2:].strip()
                lm  = re.match(r'\*\*(.+?)\*\*[:\s]*(.*)', raw)
                p   = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(5)
                ar = p.add_run('▸  ')
                ar.font.bold = True; ar.font.size = Pt(11); ar.font.color.rgb = ACCENT2_RGB
                if lm:
                    label = strip_emoji(lm.group(1))
                    rest  = strip_emoji(re.sub(r'[*`]', '', lm.group(2)))
                    lr2 = p.add_run(label + ':  ')
                    lr2.font.bold = True; lr2.font.size = Pt(10.5); lr2.font.color.rgb = ACCENT_RGB
                    if rest:
                        tr2 = p.add_run(rest); tr2.font.size = Pt(10.5); tr2.font.color.rgb = BODY_RGB
                else:
                    text = strip_emoji(re.sub(r'[*`]', '', raw))
                    tr2  = p.add_run(text); tr2.font.size = Pt(10.5); tr2.font.color.rgb = BODY_RGB
            else:
                text = strip_emoji(re.sub(r'[*`#]', '', s))
                if text:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(7)
                    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = BODY_RGB

        add_footer(doc)
        buffer = io.BytesIO()
        doc.save(buffer); buffer.seek(0)
        return StreamingResponse(buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="tinku_report.docx"'})

    except Exception as e:
        import traceback
        return JSONResponse({"error": f"DOCX failed: {str(e)}\n{traceback.format_exc()}"}, status_code=500)


# ═══════════════════════════════════════════════
# RESUME BUILDER ENDPOINT — Phase 6
# ═══════════════════════════════════════════════
from resume_builder import build_resume_pdf as _build_pdf, build_resume_docx as _build_docx

class ResumeRequest(BaseModel):
    data: dict
    format: str  # "pdf" or "docx"

@router.post("/resume")
async def build_resume(req: ResumeRequest):
    d = req.data
    name         = d.get('name', 'Your Name')
    title        = d.get('title', 'Professional')
    email        = d.get('email', '')
    phone        = d.get('phone', '')
    linkedin     = d.get('linkedin', '')
    tech_skills  = d.get('tech_skills', d.get('skills', ''))
    soft_skills  = d.get('soft_skills', '')
    experience   = d.get('experience', '')
    projects     = d.get('projects', '')
    education    = d.get('education', '')
    achievements = d.get('achievements', '')
    summary      = d.get('summary', '')
    try:
        if req.format == 'pdf':
            buf = _build_pdf(name=name, title=title, email=email, phone=phone,
                linkedin=linkedin, tech_skills=tech_skills, soft_skills=soft_skills,
                experience=experience, projects=projects, education=education,
                achievements=achievements, summary=summary)
            return StreamingResponse(buf, media_type="application/pdf",
                headers={"Content-Disposition": 'attachment; filename="resume.pdf"'})
        else:
            buf = _build_docx(name=name, title=title, email=email, phone=phone,
                linkedin=linkedin, tech_skills=tech_skills, soft_skills=soft_skills,
                experience=experience, projects=projects, education=education,
                achievements=achievements, summary=summary)
            return StreamingResponse(buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": 'attachment; filename="resume.docx"'})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
