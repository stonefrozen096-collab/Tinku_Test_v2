"""
Tinku Resume Builder — Production Ready v3
PDF + DOCX, mobile Word compatible, clean professional layout
"""
import io, re
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

_W, _H = A4
FIXED_DOC_WIDTH = _W - 1.6*cm - 1.6*cm  # 504.57 pts
A4_TWIPS     = 11906
INDENT_TWIPS = 1020

CAPS_EXACT = {'ai','ml','aws','gcp','api','apis','rest','sql','nosql','css','html',
              'js','ts','ui','ux','ios','android','ci','cd','php','xml','json',
              'jwt','http','https','tcp','ip','db','os','nlp','sdk','ide','erp',
              'crm','mba','bca','mca','seo','bi','ar','vr','iot'}

def smart_cap(text):
    if not text: return ""
    words = text.strip().split()
    result = []
    for i, word in enumerate(words):
        clean = re.sub(r'[.,;:()]', '', word)
        lower = clean.lower()
        if len(clean) > 2 and any(c.isupper() for c in clean[1:]):
            result.append(word)
        elif lower in CAPS_EXACT:
            result.append(word.replace(clean, lower.upper()))
        elif lower in ('b.tech','b.e','b.sc','m.tech','m.sc','m.e','ph.d'):
            mapping = {'b.tech':'B.Tech','m.tech':'M.Tech','b.sc':'B.Sc',
                      'm.sc':'M.Sc','ph.d':'Ph.D','b.e':'B.E','m.e':'M.E'}
            result.append(mapping.get(lower, clean.capitalize()))
        elif i == 0 and word.islower():
            result.append(word.capitalize())
        else:
            result.append(word)
    return ' '.join(result)

def clean_input(text):
    if not text: return ""
    text = str(text).strip()
    text = text.replace('&amp;','&').replace('&DS;','& DS')
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def split_items(text):
    if not text or clean_input(text).lower() in ('none','nil','na','n/a','no',''):
        return []
    return [smart_cap(i.strip()) for i in re.split(r'[,;|]', text) if i.strip()]

def skill_rows(items):
    cols = 3 if len(items) <= 6 else 4
    rows = [items[i:i+cols] for i in range(0, len(items), cols)]
    return rows, cols

def format_education(text):
    text = clean_input(text)
    text = re.sub(r'\s*&\s*', ' & ', text)
    parts = [p.strip() for p in text.split(',') if p.strip()]
    if len(parts) >= 3:
        degree  = smart_cap(parts[0])
        college = parts[1].strip()
        year    = parts[2].strip()
        if year.isdigit() and len(year) == 1:
            n   = int(year)
            sfx = 'st' if n==1 else 'nd' if n==2 else 'rd' if n==3 else 'th'
            year_str = f"{year}{sfx} Year (Pursuing)"
        elif year.isdigit() and len(year) == 4:
            year_str = f"Graduated {year}"
        else:
            year_str = year
        return degree, college, year_str
    elif len(parts) == 2:
        return smart_cap(parts[0]), parts[1], ""
    return smart_cap(text), "", ""

def parse_experience(text):
    if not text or clean_input(text).lower() in ('none','fresher','no experience','nil','na','n/a',''):
        return []
    entries = []
    for job in [j.strip() for j in re.split(r'[;\n]', text) if j.strip()]:
        parts = [p.strip() for p in job.split('|')]
        if len(parts) >= 3:
            entries.append({'company': clean_input(parts[0]),
                           'role': smart_cap(parts[1]), 'duration': parts[2]})
        elif len(parts) == 2:
            entries.append({'company': clean_input(parts[0]),
                           'role': smart_cap(parts[1]), 'duration': ''})
        else:
            entries.append({'company': '', 'role': smart_cap(job), 'duration': ''})
    return entries

def parse_projects(text):
    if not text or clean_input(text).lower() in ('none','nil','na','n/a',''):
        return []
    projects = []
    for item in [j.strip() for j in re.split(r'[;\n]', text) if j.strip()]:
        parts = [p.strip() for p in item.split('|')]
        if len(parts) >= 2:
            projects.append({'name': clean_input(parts[0]), 'desc': smart_cap(parts[1])})
        else:
            projects.append({'name': clean_input(item), 'desc': ''})
    return projects

def smart_summary(name, title, tech_skills, soft_skills, experience):
    skill_list  = split_items(tech_skills)
    soft_list   = split_items(soft_skills)
    top_skills  = ", ".join(skill_list[:3]) if skill_list else "various technologies"
    top_soft    = soft_list[0].lower() if soft_list else "collaboration"
    exp_entries = parse_experience(experience)
    has_exp     = len(exp_entries) > 0

    exp_years = 0
    if has_exp and experience:
        for m in re.finditer(r'(\d+)\s*(?:year|yr)', experience, re.IGNORECASE):
            exp_years += int(m.group(1))

    title_l = title.lower()
    if any(w in title_l for w in ['backend','back-end','api','server']):
        domain = "building scalable backend systems and APIs"
    elif any(w in title_l for w in ['frontend','front-end','ui','react','flutter']):
        domain = "crafting responsive and intuitive user interfaces"
    elif any(w in title_l for w in ['full stack','fullstack']):
        domain = "developing end-to-end web applications"
    elif any(w in title_l for w in ['data','ml','ai','machine']):
        domain = "working with data-driven and AI-powered solutions"
    elif any(w in title_l for w in ['devops','cloud','infra']):
        domain = "building and maintaining cloud infrastructure"
    elif any(w in title_l for w in ['mobile','android','ios']):
        domain = "building cross-platform mobile applications"
    else:
        domain = "building efficient and scalable software solutions"

    senior = any(w in title_l for w in ['senior','lead','principal','head','architect','manager'])
    junior = any(w in title_l for w in ['junior','associate','intern','trainee'])

    if has_exp and exp_years > 0:
        yr_str    = f"{exp_years}+ year{'s' if exp_years != 1 else ''}"
        companies = [e['company'] for e in exp_entries if e.get('company')]
        co_str    = f" across companies like {companies[0]}" if companies else ""
        if senior:
            return (f"{title} with {yr_str} of hands-on experience{co_str}, specializing in {domain}. "
                    f"Skilled in {top_skills}, with a strong track record of leading teams and delivering "
                    f"high-quality software. Known for {top_soft} and a commitment to engineering excellence.")
        else:
            return (f"{title} with {yr_str} of professional experience{co_str}, focused on {domain}. "
                    f"Proficient in {top_skills}, with a passion for writing clean, maintainable code. "
                    f"Strong {top_soft} skills with a drive to continuously learn and grow.")
    elif has_exp:
        return (f"Passionate {title} with hands-on professional experience in {domain}. "
                f"Proficient in {top_skills} and committed to building impactful, production-ready software. "
                f"Quick learner with strong {top_soft} skills and attention to detail.")
    elif junior:
        return (f"Enthusiastic {title} with a solid academic foundation in {top_skills}. "
                f"Eager to apply theoretical knowledge to real-world problems. "
                f"Strong {top_soft} skills with a passion for clean code and continuous learning.")
    else:
        return (f"Motivated {title} with a strong foundation in {top_skills}, passionate about {domain}. "
                f"Quick learner who thrives in collaborative environments and takes ownership of challenges. "
                f"Eager to contribute meaningfully while growing both technically and professionally.")


# ══════════════════════════════════════════════════════════
# PDF BUILDER — Complete Rewrite v3
# ══════════════════════════════════════════════════════════
def build_resume_pdf(name, title, email, phone, linkedin="",
                     tech_skills="", soft_skills="", experience="",
                     projects="", education="", achievements="", summary="",
                     skills=""):
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, Table, TableStyle, KeepTogether)

    if not tech_skills and skills: tech_skills = skills
    if not summary or not summary.strip():
        summary = smart_summary(name, title, tech_skills, soft_skills, experience)

    name     = clean_input(name)
    title    = clean_input(title)
    email    = clean_input(email)
    phone    = clean_input(phone)
    linkedin = clean_input(linkedin)
    if linkedin.lower() in ('none','nil','na','n/a'): linkedin = ''

    edu_deg, edu_col, edu_yr = format_education(education)
    tech_list    = split_items(tech_skills)
    soft_list    = split_items(soft_skills)
    exp_entries  = parse_experience(experience)
    proj_entries = parse_projects(projects)
    ach_list     = split_items(achievements)

    # Dense = experienced with multiple sections — use tighter spacing
    is_dense = len(exp_entries) > 0 and len(proj_entries) > 1 and len(ach_list) > 2

    # Colors
    INDIGO   = colors.HexColor('#4338CA')
    INDIGO_L = colors.HexColor('#EEF2FF')
    DARK     = colors.HexColor('#0F172A')
    TEXT     = colors.HexColor('#1E293B')
    GRAY     = colors.HexColor('#64748B')
    ACCENT   = colors.HexColor('#6366F1')
    SUBTEXT  = colors.HexColor('#94A3B8')
    WHITE    = colors.white
    BORDER   = colors.HexColor('#E2E8F0')
    LBLUE    = colors.HexColor('#C7D2FE')

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
        leftMargin=1.6*cm, rightMargin=1.6*cm,
        topMargin=0, bottomMargin=1.2*cm,
        title=f"{name} - Resume", author="Tinku AI")

    DW = FIXED_DOC_WIDTH
    story = []

    def St(n, font='Helvetica', sz=10, clr=TEXT, bold=False,
           align=TA_LEFT, before=0, after=3, lead=None):
        fn = 'Helvetica-Bold' if bold else font
        return ParagraphStyle(n, fontName=fn, fontSize=sz,
            textColor=clr, alignment=align,
            spaceBefore=before, spaceAfter=after,
            leading=lead or sz*1.35)

    # Styles
    body_after = 3 if is_dense else 6
    body_s = St('Bod', sz=11,  clr=TEXT, after=body_after, align=TA_JUSTIFY, lead=16)
    bul_after = 2 if is_dense else 5
    bul_s  = St('Bul', sz=11,  clr=TEXT, after=bul_after, lead=15)
    co_s   = St('Co',  sz=12,  clr=DARK,  bold=True, after=2)
    role_s = St('Ro',  sz=11,  clr=ACCENT, after=1)
    dur_s  = St('Du',  sz=10,  clr=GRAY,  after=4)
    pnm_s  = St('Pn',  sz=12,  clr=DARK,  bold=True, after=2)
    pds_after = 3 if is_dense else 7
    pds_s  = St('Pd',  sz=11,  clr=TEXT,  after=pds_after)
    edg_s  = St('Edg', sz=12,  clr=DARK,  bold=True, after=3)
    esb_s  = St('Esb', sz=10.5,clr=GRAY,  after=6)
    # Dynamic spacing — more space if fresher (less content) to fill page
    sec_before = 10 if is_dense else 18
    sec_s  = St('Sec', sz=11,  clr=INDIGO, bold=True, before=sec_before, after=5)
    sub_s  = St('Sub', sz=7.5, clr=SUBTEXT, align=TA_CENTER, after=0)

    def sec(text):
        story.append(KeepTogether([
            Paragraph(text.upper(), sec_s),
            HRFlowable(width="100%", thickness=1, color=INDIGO, spaceAfter=4)
        ]))

    def bul(text):
        t = smart_cap(clean_input(text))
        if t: story.append(Paragraph(f"\u2022  {t}", bul_s))

    # ── HEADER — full width, no padding runs ──
    nh = ParagraphStyle('NH', fontName='Helvetica-Bold', fontSize=26,
        textColor=WHITE, leading=26, spaceAfter=2, alignment=TA_LEFT)
    th = ParagraphStyle('TH', fontName='Helvetica', fontSize=14,
        textColor=LBLUE, leading=15, spaceAfter=2, alignment=TA_LEFT)
    ch = ParagraphStyle('CH', fontName='Helvetica', fontSize=9.5,
        textColor=colors.HexColor('#C7D2FE'), leading=12, alignment=TA_LEFT)

    cp = []
    if email:    cp.append(f"Email: {email}")
    if phone:    cp.append(f"Phone: {phone}")
    if linkedin: cp.append(f"LinkedIn: {linkedin}")

    hdr = Table([
        [Paragraph(name.upper(), nh)],
        [Paragraph(title, th)],
        [Paragraph("   |   ".join(cp), ch)],
    ], colWidths=[DW])
    hdr.setStyle(TableStyle([
        ('BACKGROUND',    (0,0),(-1,-1), INDIGO),
        ('LEFTPADDING',   (0,0),(-1,-1), 14),
        ('RIGHTPADDING',  (0,0),(-1,-1), 14),
        ('TOPPADDING',    (0,0),(0,0),   20),
        ('TOPPADDING',    (0,1),(-1,-1),  2),
        ('BOTTOMPADDING', (0,-1),(-1,-1),18),
        ('BOTTOMPADDING', (0,0),(-1,-2),  2),
        ('VALIGN',        (0,0),(-1,-1), 'MIDDLE'),
    ]))
    story.append(hdr)
    story.append(Spacer(1, 14 if not is_dense else 8))

    # ── PROFILE SUMMARY ──
    sec("Profile Summary")
    story.append(Paragraph(summary, body_s))

    # ── EDUCATION ──
    if edu_deg:
        sec("Education")
        story.append(Paragraph(edu_deg, edg_s))
        sub = []
        if edu_col: sub.append(edu_col)
        if edu_yr:  sub.append(edu_yr)
        if sub: story.append(Paragraph("  |  ".join(sub), esb_s))

    # ── TECHNICAL SKILLS — left-aligned colored tags ──
    if tech_list:
        sec("Technical Skills")
        rows, cols = skill_rows(tech_list)
        col_w = DW / cols
        tag_s = ParagraphStyle('Tag', fontName='Helvetica-Bold', fontSize=10.5,
            textColor=INDIGO, alignment=TA_LEFT, leading=14,
            leftIndent=8, rightIndent=4)
        for row in rows:
            padded = row + [''] * (cols - len(row))
            cells  = [Paragraph(s, tag_s) for s in padded]
            t = Table([cells], colWidths=[col_w]*cols, rowHeights=32)
            ts = TableStyle([
                ('TOPPADDING',    (0,0),(-1,-1), 4),
                ('BOTTOMPADDING', (0,0),(-1,-1), 4),
                ('LEFTPADDING',   (0,0),(-1,-1), 0),
                ('RIGHTPADDING',  (0,0),(-1,-1), 4),
                ('VALIGN',        (0,0),(-1,-1), 'MIDDLE'),
            ])
            for ci, skill in enumerate(padded):
                if skill:
                    ts.add('BACKGROUND', (ci,0),(ci,0), INDIGO_L)
                else:
                    ts.add('BACKGROUND', (ci,0),(ci,0), WHITE)
            t.setStyle(ts)
            story.append(t)
            story.append(Spacer(1, 5))

    # ── SOFT SKILLS ──
    if soft_list:
        sec("Soft Skills")
        story.append(Paragraph("   |   ".join(soft_list), body_s))

    # ── WORK EXPERIENCE ──
    if exp_entries:
        sec("Work Experience")
        for i, e in enumerate(exp_entries):
            if i > 0: story.append(Spacer(1, 2))
            if e['company']:
                items = [Paragraph(e['company'], co_s)]
                if e['role']:     items.append(Paragraph(e['role'], role_s))
                if e['duration']: items.append(Paragraph(e['duration'], dur_s))
                story.append(KeepTogether(items))
            else:
                bul(e['role'])

    # ── PROJECTS ──
    if proj_entries:
        sec("Projects")
        for p in proj_entries:
            items = [Paragraph(p['name'], pnm_s)]
            if p['desc']: items.append(Paragraph(p['desc'], pds_s))
            story.append(KeepTogether(items))

    # ── ACHIEVEMENTS ──
    if ach_list:
        sec("Achievements & Certifications")
        for a in ach_list: bul(a)

    # ── FOOTER ──
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceAfter=3))
    story.append(Paragraph("Resume generated by Tinku AI",
        St('Ft', font='Helvetica-Oblique', sz=7.5, clr=SUBTEXT, align=TA_CENTER)))

    doc.build(story)
    buffer.seek(0)
    return buffer


# ══════════════════════════════════════════════════════════
# DOCX BUILDER — Mobile Word Compatible v3
# ══════════════════════════════════════════════════════════
def build_resume_docx(name, title, email, phone, linkedin="",
                      tech_skills="", soft_skills="", experience="",
                      projects="", education="", achievements="", summary="",
                      skills=""):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not tech_skills and skills: tech_skills = skills
    if not summary or not summary.strip():
        summary = smart_summary(name, title, tech_skills, soft_skills, experience)

    name     = clean_input(name)
    title    = clean_input(title)
    email    = clean_input(email)
    phone    = clean_input(phone)
    linkedin = clean_input(linkedin)
    if linkedin.lower() in ('none','nil','na','n/a'): linkedin = ''

    edu_deg, edu_col, edu_yr = format_education(education)
    tech_list    = split_items(tech_skills)
    soft_list    = split_items(soft_skills)
    exp_entries  = parse_experience(experience)
    proj_entries = parse_projects(projects)
    ach_list     = split_items(achievements)
    if not ach_list or achievements.lower().strip() in ('none','nil','na','n/a'):
        ach_list = []

    # Dense layout for experienced resumes
    is_dense_doc = len(exp_entries) > 0 and len(proj_entries) > 1

    # Colors
    C_IND  = RGBColor(0x43,0x38,0xCA)
    C_SKL  = RGBColor(0x58,0x50,0xE8)   # slightly lighter — skills vs headings
    C_ACC  = RGBColor(0x63,0x66,0xF1)
    C_DARK = RGBColor(0x0F,0x17,0x2A)
    C_TEXT = RGBColor(0x1E,0x29,0x3B)
    C_GRAY = RGBColor(0x64,0x74,0x8B)
    C_SUB  = RGBColor(0x94,0xA3,0xB8)
    C_WHT  = RGBColor(0xFF,0xFF,0xFF)
    C_LBLU = RGBColor(0xA5,0xB4,0xFC)
    C_SEP  = RGBColor(0xC7,0xD2,0xFE)
    C_BDR  = RGBColor(0xE2,0xE8,0xF0)

    doc = Document()
    for s in doc.sections:
        s.top_margin    = Twips(0)
        s.bottom_margin = Cm(1.5)
        s.left_margin   = Twips(0)
        s.right_margin  = Twips(0)

    # ── XML helpers ──
    def pPr(p): return p._p.get_or_add_pPr()

    def shading(p, fill):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), fill); pPr(p).append(shd)

    def indent(p, left=INDENT_TWIPS, right=INDENT_TWIPS):
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(left))
        ind.set(qn('w:right'), str(right))
        pPr(p).append(ind)

    def spacing(p, before=0, after=60):
        pr = pPr(p)
        spc = OxmlElement('w:spacing')
        spc.set(qn('w:before'), str(int(before)))
        spc.set(qn('w:after'),  str(int(after)))
        spc.set(qn('w:line'), '276')
        spc.set(qn('w:lineRule'), 'auto')
        pr.append(spc)

    def border_bottom(p, color='4338CA', sz='8'):
        pr = pPr(p)
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'), sz)
        bot.set(qn('w:space'),'1');    bot.set(qn('w:color'), color)
        pBdr.append(bot); pr.append(pBdr)

    def run(p, text, sz, rgb, bold=False, italic=False):
        r = p.add_run(text)
        r.font.size=Pt(sz); r.font.bold=bold; r.font.italic=italic
        r.font.color.rgb=rgb
        return r

    # ── Header lines (paragraph shading = full width on all Word versions) ──
    def hline(text, sz, rgb, bold=False, before=0, after=50):
        p = doc.add_paragraph()
        shading(p, '4338CA')
        indent(p, left=INDENT_TWIPS, right=INDENT_TWIPS)
        spacing(p, before=before, after=after)
        run(p, text, sz, rgb, bold=bold)

    hline(name.upper(), 24, C_WHT,  bold=True, before=260, after=70)
    hline(title,        13, C_LBLU, before=0,  after=70)
    cp = []
    if email:    cp.append(f"Email: {email}")
    if phone:    cp.append(f"Phone: {phone}")
    if linkedin: cp.append(f"LinkedIn: {linkedin}")
    hline("   |   ".join(cp), 9.5, C_SEP, before=0, after=320 if not is_dense_doc else 260)

    # ── Section heading ──
    def heading(label):
        label = re.sub(r'(\w)&(\w)', r'\1 & \2', label)
        p = doc.add_paragraph()
        indent(p)
        sp_before = 140 if is_dense_doc else 360
        spacing(p, before=sp_before, after=80)
        border_bottom(p)
        run(p, label.upper(), 11.5, C_IND, bold=True)

    # ── Body paragraph ──
    _body_after = 50 if (len(exp_entries) > 0 and len(proj_entries) > 1) else 130
    def body(text, sz=11, rgb=None, bold=False, italic=False, before=0, after=None):
        after = after if after is not None else _body_after
        p = doc.add_paragraph()
        indent(p)
        spacing(p, before=before, after=after)
        run(p, smart_cap(clean_input(text)), sz, rgb or C_TEXT, bold=bold, italic=italic)
        return p

    # ── Bullet ──
    def bullet(text):
        p = doc.add_paragraph()
        indent(p, left=INDENT_TWIPS+100, right=INDENT_TWIPS)
        spacing(p, before=0, after=40)
        run(p, f"\u2022  {smart_cap(clean_input(text))}", 10, C_TEXT)

    # ── Skills row ──
    def skills_row(items):
        p = doc.add_paragraph()
        indent(p)
        spacing(p, before=0, after=40)
        for i, skill in enumerate(items):
            if i > 0:
                sep = p.add_run("   |   ")
                sep.font.size=Pt(10)
                sep.font.color.rgb=C_SEP
            r2 = p.add_run(skill)
            r2.font.size=Pt(11); r2.font.bold=False
            r2.font.color.rgb=C_SKL

    # ── Thin divider ──
    def thin_divider():
        p = doc.add_paragraph()
        indent(p)
        spacing(p, before=0, after=30)
        border_bottom(p, color='E2E8F0', sz='2')

    # ── SECTIONS ──
    heading("Profile Summary")
    body(summary, after=140 if not is_dense_doc else 80)

    if edu_deg:
        heading("Education")
        body(edu_deg, sz=12, rgb=C_DARK, bold=True, after=35)
        sub = []
        if edu_col: sub.append(edu_col)
        if edu_yr:  sub.append(edu_yr)
        if sub: body("   |   ".join(sub), sz=10, rgb=C_GRAY, after=100 if not is_dense_doc else 55)

    if tech_list:
        heading("Technical Skills")
        rows, cols = skill_rows(tech_list)
        for row in rows:
            skills_row(row)

    if soft_list:
        heading("Soft Skills")
        body("   |   ".join(soft_list), after=130 if not is_dense_doc else 70)

    if exp_entries:
        heading("Work Experience")
        for i, e in enumerate(exp_entries):
            if i > 0: thin_divider()
            if e['company']:
                body(e['company'], sz=12, rgb=C_DARK, bold=True, after=35)
                if e['role']:     body(e['role'],     sz=11, rgb=C_ACC,  after=25)
                if e['duration']: body(e['duration'], sz=10, rgb=C_GRAY, italic=True, after=65)
            else:
                bullet(e['role'])

    if proj_entries:
        heading("Projects")
        for p in proj_entries:
            body(p['name'], sz=12, rgb=C_DARK, bold=True, after=30)
            if p['desc']: body(p['desc'], sz=11, after=100 if not is_dense_doc else 60)

    if ach_list:
        heading("Achievements & Certifications")
        for a in ach_list: bullet(a)

    # Footer
    fp = doc.add_paragraph()
    indent(fp)
    spacing(fp, before=160, after=0)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = fp.add_run("Resume generated by Tinku AI")
    r2.font.size=Pt(8); r2.font.italic=True; r2.font.color.rgb=C_SUB

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf
